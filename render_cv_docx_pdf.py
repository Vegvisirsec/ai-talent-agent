import argparse
import re
import tempfile
from pathlib import Path


IMAGE_RE = re.compile(r"!\[([^\]]*)\]\(([^)]+)\)")
BOLD_RE = re.compile(r"\*\*([^*]+)\*\*")
PLACEHOLDER_IMAGE_RE = re.compile(
    r"!\[[^\]]*Profile Photo[^\]]*\]\((photo-placeholder\.jpg|portrait\.png)\)",
    re.IGNORECASE,
)
DEFAULT_MAX_BYTES = 1_572_864
IMAGE_COMPRESSION_STEPS = (
    (1200, 82),
    (1000, 72),
    (900, 64),
    (800, 56),
    (700, 48),
    (600, 40),
    (500, 34),
)


def resolve_cli_path(value: str | None, base_dir: Path | None = None) -> Path | None:
    if not value:
        return None

    candidate = Path(value).expanduser()
    if candidate.is_absolute():
        return candidate.resolve()

    if base_dir is not None:
        return (base_dir / candidate).resolve()

    return candidate.resolve()


def strip_markdown(text: str) -> str:
    text = IMAGE_RE.sub("", text)
    text = BOLD_RE.sub(r"\1", text)
    text = re.sub(r"\[([^\]]+)\]\(([^)]+)\)", r"\1 (\2)", text)
    text = re.sub(r"`([^`]+)`", r"\1", text)
    return text.strip()


def parse_markdown(markdown_text: str):
    blocks = []
    current_list = []

    def flush_list():
        nonlocal current_list
        if current_list:
            blocks.append(("list", current_list))
            current_list = []

    for raw_line in markdown_text.splitlines():
        line = raw_line.rstrip()
        stripped = line.strip()

        if not stripped:
            flush_list()
            continue

        if stripped == "---":
            flush_list()
            blocks.append(("rule", None))
            continue

        if stripped.startswith("# "):
            flush_list()
            blocks.append(("h1", strip_markdown(stripped[2:])))
            continue

        if stripped.startswith("## "):
            flush_list()
            blocks.append(("h2", strip_markdown(stripped[3:])))
            continue

        if stripped.startswith("### "):
            flush_list()
            blocks.append(("h3", strip_markdown(stripped[4:])))
            continue

        image_match = IMAGE_RE.fullmatch(stripped)
        if image_match:
            flush_list()
            blocks.append(("image", image_match.group(2)))
            continue

        if stripped.startswith("- "):
            current_list.append(strip_markdown(stripped[2:]))
            continue

        blocks.append(("p", strip_markdown(stripped)))

    flush_list()
    return blocks


def prepare_markdown(markdown_text: str, base_dir: Path, portrait_path: Path | None = None):
    portrait_path = portrait_path or (base_dir / "portrait.png")
    embedded_portrait = False
    portrait_reference = None

    if portrait_path == (base_dir / "portrait.png"):
        portrait_reference = "portrait.png"
    else:
        portrait_reference = portrait_path.resolve().as_posix()

    if portrait_path.exists():
        if PLACEHOLDER_IMAGE_RE.search(markdown_text):
            markdown_text = PLACEHOLDER_IMAGE_RE.sub(
                f"![Profile Photo]({portrait_reference})", markdown_text, count=1
            )
            embedded_portrait = True
        elif not IMAGE_RE.search(markdown_text):
            lines = markdown_text.splitlines()
            if lines and lines[0].startswith("# "):
                lines.insert(1, "")
                lines.insert(2, f"![Profile Photo]({portrait_reference})")
                lines.insert(3, "")
                markdown_text = "\n".join(lines)
            else:
                markdown_text = f"![Profile Photo]({portrait_reference})\n\n" + markdown_text
            embedded_portrait = True

    return markdown_text, portrait_path, portrait_reference, embedded_portrait


def resolve_image_path(base_dir: Path, value: str) -> Path:
    image_path = Path(value)
    if image_path.is_absolute():
        return image_path
    return (base_dir / value).resolve()


def replace_image_blocks(blocks, replacement_path: Path):
    updated = []
    for kind, value in blocks:
        if kind == "image":
            updated.append((kind, str(replacement_path)))
        else:
            updated.append((kind, value))
    return updated


def build_compressed_image(source_path: Path, output_path: Path, max_edge: int, quality: int):
    try:
        from PIL import Image
    except ImportError:
        return False

    with Image.open(source_path) as image:
        image = image.convert("RGBA")
        background = Image.new("RGB", image.size, "white")
        background.paste(image, mask=image.getchannel("A"))
        image = background
        image.thumbnail((max_edge, max_edge))
        image.save(output_path, format="JPEG", quality=quality, optimize=True, progressive=True)
    return True


def render_with_size_cap(render_fn, blocks, output_path: Path, base_dir: Path, theme: str, max_bytes: int):
    image_blocks = [value for kind, value in blocks if kind == "image"]
    render_fn(blocks, output_path, base_dir, theme)

    if output_path.stat().st_size <= max_bytes or not image_blocks:
        return output_path.stat().st_size, False

    original_image_path = resolve_image_path(base_dir, image_blocks[0])
    if not original_image_path.exists():
        return output_path.stat().st_size, False

    with tempfile.TemporaryDirectory() as temp_dir:
        temp_dir_path = Path(temp_dir)
        compressed_any = False
        for idx, (max_edge, quality) in enumerate(IMAGE_COMPRESSION_STEPS, start=1):
            compressed_path = temp_dir_path / f"portrait_{idx}.jpg"
            if not build_compressed_image(original_image_path, compressed_path, max_edge, quality):
                break
            compressed_any = True
            candidate_blocks = replace_image_blocks(blocks, compressed_path)
            render_fn(candidate_blocks, output_path, base_dir, theme)
            if output_path.stat().st_size <= max_bytes:
                return output_path.stat().st_size, True

    return output_path.stat().st_size, compressed_any


def render_docx(blocks, output_path: Path, base_dir: Path, theme: str):
    try:
        from docx import Document
        from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
        from docx.shared import Inches, Pt, RGBColor
    except ImportError as exc:
        raise RuntimeError(
            "DOCX export requires python-docx. Install it with: pip install python-docx"
        ) from exc

    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(0.6)
    section.bottom_margin = Inches(0.6)
    section.left_margin = Inches(0.7)
    section.right_margin = Inches(0.7)

    normal = document.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)
    heading_color = RGBColor(0x10, 0x20, 0x33)
    body_color = RGBColor(0x18, 0x21, 0x2B)

    for kind, value in blocks:
        if kind == "h1":
            p = document.add_paragraph()
            p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            run = p.add_run(value)
            run.bold = True
            run.font.size = Pt(18)
            run.font.color.rgb = heading_color
        elif kind == "h2":
            p = document.add_paragraph()
            run = p.add_run(value)
            run.bold = True
            run.font.size = Pt(12)
            run.font.color.rgb = heading_color
        elif kind == "h3":
            p = document.add_paragraph()
            run = p.add_run(value)
            run.bold = True
            run.font.size = Pt(11)
            run.font.color.rgb = heading_color
        elif kind == "image":
            image_path = resolve_image_path(base_dir, value)
            if image_path.exists():
                document.add_picture(str(image_path), width=Inches(1.35))
        elif kind == "p":
            p = document.add_paragraph(value)
            for run in p.runs:
                run.font.color.rgb = body_color
        elif kind == "list":
            for item in value:
                p = document.add_paragraph(item, style="List Bullet")
                for run in p.runs:
                    run.font.color.rgb = body_color
        elif kind == "rule":
            document.add_paragraph("")

    document.save(output_path)


def render_pdf(blocks, output_path: Path, base_dir: Path, theme: str):
    try:
        from reportlab.lib import colors
        from reportlab.lib.enums import TA_LEFT
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
        from reportlab.lib.units import mm
        from reportlab.platypus import (
            Image,
            ListFlowable,
            ListItem,
            Paragraph,
            SimpleDocTemplate,
            Spacer,
        )
    except ImportError as exc:
        raise RuntimeError(
            "PDF export requires reportlab. Install it with: pip install reportlab"
        ) from exc

    doc = SimpleDocTemplate(
        str(output_path),
        pagesize=A4,
        leftMargin=18 * mm,
        rightMargin=18 * mm,
        topMargin=14 * mm,
        bottomMargin=14 * mm,
    )

    palette = {
        "light": {
            "page_bg": colors.white,
            "body": colors.HexColor("#18212B"),
            "heading": colors.HexColor("#102033"),
            "muted": colors.HexColor("#5B6878"),
        },
        "dark": {
            "page_bg": colors.HexColor("#161B22"),
            "body": colors.HexColor("#E6EDF3"),
            "heading": colors.HexColor("#F0F6FC"),
            "muted": colors.HexColor("#9FB0C3"),
        },
    }[theme]

    styles = getSampleStyleSheet()
    body = ParagraphStyle(
        "CVBody",
        parent=styles["BodyText"],
        fontName="Helvetica",
        fontSize=10,
        leading=13,
        alignment=TA_LEFT,
        spaceAfter=6,
        textColor=palette["body"],
    )
    h1 = ParagraphStyle(
        "CVH1",
        parent=body,
        fontName="Helvetica-Bold",
        fontSize=18,
        leading=22,
        spaceAfter=10,
        textColor=palette["heading"],
    )
    h2 = ParagraphStyle(
        "CVH2",
        parent=body,
        fontName="Helvetica-Bold",
        fontSize=12,
        leading=15,
        textColor=palette["heading"],
        spaceBefore=10,
        spaceAfter=6,
    )
    h3 = ParagraphStyle(
        "CVH3",
        parent=body,
        fontName="Helvetica-Bold",
        fontSize=10.5,
        leading=13,
        spaceBefore=6,
        spaceAfter=2,
        textColor=palette["heading"],
    )

    story = []
    for kind, value in blocks:
        if kind == "h1":
            story.append(Paragraph(value, h1))
        elif kind == "h2":
            story.append(Paragraph(value, h2))
        elif kind == "h3":
            story.append(Paragraph(value, h3))
        elif kind == "image":
            image_path = resolve_image_path(base_dir, value)
            if image_path.exists():
                image = Image(str(image_path), width=34 * mm, height=42.5 * mm, kind="proportional")
                story.append(image)
                story.append(Spacer(1, 6))
        elif kind == "p":
            story.append(Paragraph(
                value.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;"),
                body,
            ))
        elif kind == "list":
            items = [
                ListItem(Paragraph(item.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;"), body))
                for item in value
            ]
            story.append(ListFlowable(items, bulletType="bullet", leftIndent=14))
            story.append(Spacer(1, 4))
        elif kind == "rule":
            story.append(Spacer(1, 6))

    def draw_page_background(canvas, _doc):
        if theme != "dark":
            return
        canvas.saveState()
        canvas.setFillColor(palette["page_bg"])
        canvas.rect(0, 0, A4[0], A4[1], fill=1, stroke=0)
        canvas.restoreState()

    if theme == "dark":
        doc.build(
            story,
            onFirstPage=draw_page_background,
            onLaterPages=draw_page_background,
        )
    else:
        doc.build(story)


def main():
    parser = argparse.ArgumentParser(description="Render a CV markdown file to DOCX and/or PDF.")
    parser.add_argument("input", nargs="?", help="Path to the markdown CV file")
    parser.add_argument("--path", dest="input_path", help="Path to the markdown CV file")
    parser.add_argument("--docx", action="store_true", help="Generate a .docx file")
    parser.add_argument("--pdf", action="store_true", help="Generate a .pdf file")
    parser.add_argument("--both", action="store_true", help="Generate both .docx and .pdf")
    parser.add_argument(
        "--theme",
        choices=["light", "dark"],
        default="light",
        help="Visual theme for rendered outputs. PDF supports both themes well; DOCX stays conservative for compatibility.",
    )
    parser.add_argument(
        "--max-bytes",
        type=int,
        default=DEFAULT_MAX_BYTES,
        help="Best-effort maximum output size in bytes for DOCX/PDF. Defaults to 1572864 (1.5 MB).",
    )
    parser.add_argument(
        "--portrait",
        help="Optional path to a portrait image. Defaults to portrait.png next to the markdown file.",
    )
    args = parser.parse_args()

    selected_input = args.input_path or args.input
    if not selected_input:
        parser.error("Provide a markdown CV path either as a positional argument or with --path.")

    if not (args.docx or args.pdf or args.both):
        parser.error("Specify at least one output format: --docx, --pdf, or --both.")

    input_path = Path(selected_input)
    portrait_path = resolve_cli_path(args.portrait)
    markdown_text = input_path.read_text(encoding="utf-8")
    markdown_text, portrait_path, portrait_reference, embedded_portrait = prepare_markdown(
        markdown_text, input_path.parent, portrait_path
    )
    blocks = parse_markdown(markdown_text)

    generate_docx = args.docx or args.both
    generate_pdf = args.pdf or args.both

    if generate_docx:
        docx_path = input_path.with_suffix(".docx")
        docx_size, docx_compressed = render_with_size_cap(
            render_docx, blocks, docx_path, input_path.parent, args.theme, args.max_bytes
        )
        print(f"Rendered {input_path} -> {docx_path}")
        if docx_compressed:
            print("Applied portrait compression to reduce DOCX size.")
        if docx_size > args.max_bytes:
            print(
                f"Warning: {docx_path.name} is {docx_size} bytes, which exceeds the target of {args.max_bytes} bytes."
            )

    if generate_pdf:
        pdf_path = input_path.with_suffix(".pdf")
        pdf_size, pdf_compressed = render_with_size_cap(
            render_pdf, blocks, pdf_path, input_path.parent, args.theme, args.max_bytes
        )
        print(f"Rendered {input_path} -> {pdf_path}")
        if pdf_compressed:
            print("Applied portrait compression to reduce PDF size.")
        if pdf_size > args.max_bytes:
            print(
                f"Warning: {pdf_path.name} is {pdf_size} bytes, which exceeds the target of {args.max_bytes} bytes."
            )

    if embedded_portrait:
        print(f"Embedded portrait from {portrait_path}")
    elif PLACEHOLDER_IMAGE_RE.search(markdown_text):
        print(f"No portrait embedded. Expected file at {portrait_path}")


if __name__ == "__main__":
    main()
